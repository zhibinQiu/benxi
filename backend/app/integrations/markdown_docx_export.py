"""Markdown 研究报告导出为 Word (.docx)。"""

from __future__ import annotations

import base64
import io
import logging
import re
import zlib
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument

logger = logging.getLogger(__name__)

_CN_FONT = "宋体"
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_ORDERED_RE = re.compile(r"^(\d+)\.\s+(.+)$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.+)$")
_TABLE_SEP_RE = re.compile(r"^\|?[\s:-]+\|[\s|:-]+$")
_INLINE_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_INLINE_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_IMAGE_MD_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_FENCE_OPEN_RE = re.compile(r"^```(\w*)\s*$")
_FENCE_CLOSE_RE = re.compile(r"^```\s*$")
_DATA_URI_RE = re.compile(r"^data:image/[^;]+;base64,(.+)$", re.I)
_EXPORT_CITATION_RE = re.compile(r"[\[【]\d{1,2}[\]】]")
_REPORT_HEADING_RE = re.compile(r"^#{1,3}\s+\S")
_PREAMBLE_LINE_RE = re.compile(
    r"(?:"
    r"^好的[，,]|遵照.{0,6}指示|我将以|为您撰写|严格遵循|深度融合|"
    r"首轮正文|确保内容详实|将以.{0,8}身份|报告将严格|力求达到|"
    r"^以下是|^以下为|^接下来"
    r")",
    re.I,
)


def _safe_filename_stem(title: str, *, fallback: str = "研究报告") -> str:
    stem = re.sub(r'[\\/:*?"<>|]+', "_", (title or "").strip())
    stem = stem.strip("._ ") or fallback
    return stem[:80]


def _set_cn_normal_font(doc: DocxDocument, *, font_name: str = _CN_FONT) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _configure_heading_styles(doc: DocxDocument, *, font_name: str = _CN_FONT) -> None:
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = font_name
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _apply_run_font(run, *, font_name: str = _CN_FONT, bold: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    run.font.name = font_name
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if bold is not None:
        run.bold = bold


def _strip_inline_md(text: str) -> str:
    text = _INLINE_BOLD_RE.sub(r"\1", text)
    text = _INLINE_ITALIC_RE.sub(r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def _add_rich_paragraph(doc: DocxDocument, text: str, *, style: str | None = None) -> None:
    from docx.shared import Pt

    clean = _strip_inline_md(text)
    if not clean:
        doc.add_paragraph("")
        return
    para = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold_match = _INLINE_BOLD_RE.fullmatch(part)
        if bold_match:
            run = para.add_run(_strip_inline_md(bold_match.group(1)))
            _apply_run_font(run, bold=True)
            run.font.size = Pt(11)
            continue
        run = para.add_run(_strip_inline_md(part))
        _apply_run_font(run)
        run.font.size = Pt(11)


def _parse_table_row(line: str) -> list[str]:
    raw = line.strip().strip("|")
    return [_strip_inline_md(cell.strip()) for cell in raw.split("|")]


def _add_table(doc: DocxDocument, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    if cols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            value = row[c_idx] if c_idx < len(row) else ""
            table.rows[r_idx].cells[c_idx].text = value


def _decode_image_src(src: str) -> bytes | None:
    src = (src or "").strip()
    if not src:
        return None
    data_match = _DATA_URI_RE.match(src)
    if data_match:
        try:
            return base64.b64decode(data_match.group(1))
        except Exception:
            return None
    if src.startswith("http://") or src.startswith("https://"):
        try:
            import httpx

            resp = httpx.get(src, timeout=30, follow_redirects=True)
            resp.raise_for_status()
            return resp.content
        except Exception as exc:
            logger.warning("Word 导出图片下载失败: %s", exc)
            return None
    return None


def _add_image_from_src(doc: DocxDocument, src: str, *, alt: str = "") -> bool:
    from docx.shared import Inches

    data = _decode_image_src(src)
    if not data:
        if alt:
            _add_rich_paragraph(doc, alt)
        return False
    try:
        stream = io.BytesIO(data)
        para = doc.add_paragraph()
        para.alignment = 1  # center
        run = para.add_run()
        run.add_picture(stream, width=Inches(5.8))
        return True
    except Exception as exc:
        logger.warning("Word 导出嵌入图片失败: %s", exc)
        if alt:
            _add_rich_paragraph(doc, alt)
        return False


def _is_valid_png(data: bytes | None) -> bool:
    return bool(data) and data[:8] == b"\x89PNG\r\n\x1a\n"


def _render_mermaid_png_kroki(text: str) -> bytes | None:
    try:
        import httpx

        resp = httpx.post(
            "https://kroki.io/mermaid/png",
            content=text.encode("utf-8"),
            headers={"Content-Type": "text/plain"},
            timeout=45,
        )
        resp.raise_for_status()
        if _is_valid_png(resp.content):
            return resp.content
    except Exception as exc:
        logger.warning("Mermaid Kroki 渲染失败: %s", exc)
    return None


def _render_mermaid_png_mermaid_ink(text: str) -> bytes | None:
    try:
        compressed = zlib.compress(text.encode("utf-8"), 9)
        encoded = base64.urlsafe_b64encode(compressed).decode("ascii")
        import httpx

        resp = httpx.get(
            f"https://mermaid.ink/img/pako:{encoded}",
            timeout=45,
            follow_redirects=True,
        )
        resp.raise_for_status()
        if _is_valid_png(resp.content):
            return resp.content
    except Exception as exc:
        logger.warning("Mermaid.ink 渲染失败: %s", exc)
    return None


def _render_mermaid_png(source: str) -> bytes | None:
    from app.core.report_mermaid_sanitize import sanitize_mermaid_source

    text = sanitize_mermaid_source((source or "").strip())
    if not text:
        return None
    png = _render_mermaid_png_kroki(text)
    if png:
        return png
    return _render_mermaid_png_mermaid_ink(text)


def strip_export_citation_markers(text: str) -> str:
    """Word 导出：移除正文句末引用编号 [1]、【2】等。"""
    lines: list[str] = []
    for line in (text or "").splitlines():
        cleaned = _EXPORT_CITATION_RE.sub("", line)
        cleaned = re.sub(r"[ \t]{2,}", " ", cleaned).rstrip()
        lines.append(cleaned)
    return "\n".join(lines).strip()


def strip_report_generation_preamble(text: str) -> str:
    """移除报告正文前的寒暄/任务复述，从首个标题或实质段落开始。"""
    raw = (text or "").strip()
    if not raw:
        return raw

    lines = raw.splitlines()
    for idx, line in enumerate(lines):
        if _REPORT_HEADING_RE.match(line.strip()):
            if idx == 0:
                return raw
            return "\n".join(lines[idx:]).strip()

    kept: list[str] = []
    skipping = True
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if not skipping:
                kept.append(line)
            continue
        if skipping and _PREAMBLE_LINE_RE.search(stripped):
            continue
        skipping = False
        kept.append(line)
    return "\n".join(kept).strip() or raw


def prepare_report_markdown_for_export(markdown_text: str) -> str:
    """Word/入库导出：仅保留报告正文，去除开场白、引用编号与来源章节。"""
    from app.services.knowledge_qa.citations import strip_answer_source_narrative

    text = strip_report_generation_preamble(markdown_text)
    text = strip_answer_source_narrative(text)
    text = strip_export_citation_markers(text)
    return text.strip()


def _add_mermaid_block(doc: DocxDocument, source: str) -> None:
    png = _render_mermaid_png(source)
    if png:
        try:
            from docx.shared import Inches

            stream = io.BytesIO(png)
            para = doc.add_paragraph()
            para.alignment = 1
            run = para.add_run()
            run.add_picture(stream, width=Inches(5.8))
            return
        except Exception as exc:
            logger.warning("Word 导出嵌入 Mermaid PNG 失败: %s", exc)
    _add_rich_paragraph(doc, "[图表渲染失败]")


def _add_heading(doc: DocxDocument, text: str, *, level: int) -> None:
    from docx.shared import Pt

    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _apply_run_font(run, bold=True)
        if level == 0:
            run.font.size = Pt(18)


def markdown_to_docx_bytes(
    *,
    title: str,
    markdown_text: str,
    for_export: bool = False,
) -> bytes:
    from docx import Document

    body = prepare_report_markdown_for_export(markdown_text) if for_export else (markdown_text or "")
    doc = Document()
    _set_cn_normal_font(doc)
    _configure_heading_styles(doc)

    report_title = (title or "研究报告").strip() or "研究报告"
    _add_heading(doc, report_title, level=0)

    lines = body.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        fence_open = _FENCE_OPEN_RE.match(stripped)
        if fence_open:
            lang = (fence_open.group(1) or "").lower()
            idx += 1
            block_lines: list[str] = []
            while idx < len(lines) and not _FENCE_CLOSE_RE.match(lines[idx].strip()):
                block_lines.append(lines[idx].rstrip())
                idx += 1
            if idx < len(lines):
                idx += 1
            body = "\n".join(block_lines)
            if lang == "mermaid":
                _add_mermaid_block(doc, body)
            elif body.strip():
                _add_rich_paragraph(doc, body)
            continue

        image = _IMAGE_MD_RE.fullmatch(stripped)
        if image:
            _add_image_from_src(doc, image.group(2), alt=image.group(1))
            idx += 1
            continue

        idx += 1

        if not line.strip():
            continue
        if line.strip() in {"---", "***", "___"}:
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            level = min(len(heading.group(1)), 4)
            text = _strip_inline_md(heading.group(2))
            _add_heading(doc, text, level=level)
            continue

        if line.startswith(">"):
            _add_rich_paragraph(doc, line.lstrip(">").strip())
            continue

        ordered = _ORDERED_RE.match(line)
        if ordered:
            _add_rich_paragraph(doc, f"{ordered.group(1)}. {ordered.group(2)}")
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(_strip_inline_md(bullet.group(1)))
            _apply_run_font(run)
            continue

        if "|" in line and idx < len(lines) and _TABLE_SEP_RE.match(lines[idx].strip()):
            table_rows = [_parse_table_row(line)]
            idx += 1
            while idx < len(lines) and "|" in lines[idx]:
                if not lines[idx].strip():
                    break
                table_rows.append(_parse_table_row(lines[idx]))
                idx += 1
            _add_table(doc, table_rows)
            continue

        _add_rich_paragraph(doc, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx_download_filename(title: str) -> str:
    return f"{_safe_filename_stem(title)}.docx"
