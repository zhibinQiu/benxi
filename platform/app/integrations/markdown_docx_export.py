"""Markdown 研究报告导出为 Word (.docx)。"""

from __future__ import annotations

import io
import re
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_ORDERED_RE = re.compile(r"^(\d+)\.\s+(.+)$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.+)$")
_TABLE_SEP_RE = re.compile(r"^\|?[\s:-]+\|[\s|:-]+$")
_INLINE_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_INLINE_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _safe_filename_stem(title: str, *, fallback: str = "研究报告") -> str:
    stem = re.sub(r'[\\/:*?"<>|]+', "_", (title or "").strip())
    stem = stem.strip("._ ") or fallback
    return stem[:80]


def _set_cn_normal_font(doc: DocxDocument, *, font_name: str = "微软雅黑") -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _apply_run_font(run, *, font_name: str = "微软雅黑") -> None:
    from docx.oxml.ns import qn

    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _strip_inline_md(text: str) -> str:
    text = _INLINE_BOLD_RE.sub(r"\1", text)
    text = _INLINE_ITALIC_RE.sub(r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def _add_rich_paragraph(doc: DocxDocument, text: str, *, style: str | None = None) -> None:
    from docx.shared import Pt

    clean = _strip_inline_md(text)
    if not clean:
        doc.add_paragraph("")
        return
    para = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold_match = _INLINE_BOLD_RE.fullmatch(part)
        if bold_match:
            run = para.add_run(_strip_inline_md(bold_match.group(1)))
            run.bold = True
            _apply_run_font(run)
            run.font.size = Pt(11)
            continue
        run = para.add_run(_strip_inline_md(part))
        _apply_run_font(run)
        run.font.size = Pt(11)


def _parse_table_row(line: str) -> list[str]:
    raw = line.strip().strip("|")
    return [_strip_inline_md(cell.strip()) for cell in raw.split("|")]


def _add_table(doc: DocxDocument, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    if cols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            value = row[c_idx] if c_idx < len(row) else ""
            table.rows[r_idx].cells[c_idx].text = value


def markdown_to_docx_bytes(*, title: str, markdown_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    _set_cn_normal_font(doc)

    report_title = (title or "研究报告").strip() or "研究报告"
    title_para = doc.add_heading(report_title, level=0)
    for run in title_para.runs:
        _apply_run_font(run)
        run.font.size = Pt(18)

    lines = (markdown_text or "").splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        idx += 1

        if not line.strip():
            continue
        if line.strip() in {"---", "***", "___"}:
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            level = min(len(heading.group(1)), 4)
            text = _strip_inline_md(heading.group(2))
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                _apply_run_font(run)
            continue

        if line.startswith(">"):
            _add_rich_paragraph(doc, line.lstrip(">").strip())
            continue

        ordered = _ORDERED_RE.match(line)
        if ordered:
            _add_rich_paragraph(doc, f"{ordered.group(1)}. {ordered.group(2)}")
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(_strip_inline_md(bullet.group(1)))
            _apply_run_font(run)
            continue

        if "|" in line and idx < len(lines) and _TABLE_SEP_RE.match(lines[idx].strip()):
            table_rows = [_parse_table_row(line)]
            idx += 1  # skip separator
            while idx < len(lines) and "|" in lines[idx]:
                if not lines[idx].strip():
                    break
                table_rows.append(_parse_table_row(lines[idx]))
                idx += 1
            _add_table(doc, table_rows)
            continue

        _add_rich_paragraph(doc, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx_download_filename(title: str) -> str:
    return f"{_safe_filename_stem(title)}.docx"
