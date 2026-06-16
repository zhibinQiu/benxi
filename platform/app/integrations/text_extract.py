"""CPU-only document text extraction (no KnowFlow / GPU)."""

from __future__ import annotations

import io
import re
import uuid
import zipfile
from dataclasses import dataclass, field
from html import unescape
from xml.etree import ElementTree

from app.integrations.html_markdown import html_to_markdown


@dataclass
class TextBlock:
    text: str
    page: int = 1
    bbox: list[float] | None = None


@dataclass
class ParsedDocument:
    document_id: uuid.UUID
    file_name: str
    full_text: str
    pages: list[dict] = field(default_factory=list)
    parse_quality: str = "text_layer"
    warning: str | None = None


_PLAIN_TEXT_EXTENSIONS = (
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".log",
    ".json",
    ".xml",
    ".yaml",
    ".yml",
    ".ini",
    ".conf",
    ".properties",
)


def _mime_base(mime_type: str) -> str:
    return (mime_type or "").split(";")[0].strip().lower()


def _text_pages(full_text: str, *, paragraphs: list[str] | None = None) -> list[dict]:
    paras = paragraphs or split_paragraphs(full_text)
    if not paras and full_text.strip():
        paras = [full_text.strip()]
    return [
        {
            "page": 1,
            "text": full_text,
            "blocks": [{"text": t, "bbox": None} for t in paras],
        }
    ]


def _parsed_text_doc(
    *,
    document_id: uuid.UUID,
    file_name: str,
    full_text: str,
    parse_quality: str = "text_layer",
    warning: str | None = None,
    paragraphs: list[str] | None = None,
) -> ParsedDocument:
    text = (full_text or "").strip()
    return ParsedDocument(
        document_id=document_id,
        file_name=file_name,
        full_text=text,
        pages=_text_pages(text, paragraphs=paragraphs),
        parse_quality=parse_quality,
        warning=warning,
    )


def extract_text_from_bytes(
    data: bytes,
    *,
    document_id: uuid.UUID,
    file_name: str,
    mime_type: str = "",
) -> ParsedDocument:
    lower = (file_name or "").lower()
    mime = _mime_base(mime_type)

    if lower.endswith(".pdf") or mime == "application/pdf":
        return _extract_pdf(data, document_id=document_id, file_name=file_name)
    if lower.endswith((".doc", ".docx", ".dot", ".dotx")) or "word" in mime:
        return _extract_docx(data, document_id=document_id, file_name=file_name)
    if lower.endswith((".html", ".htm")) or "html" in mime:
        return _extract_html(data, document_id=document_id, file_name=file_name)
    if lower.endswith((".xlsx", ".xlsm")) or mime in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ) or "spreadsheet" in mime or "excel" in mime:
        return _extract_excel(data, document_id=document_id, file_name=file_name)
    if lower.endswith(".xls"):
        return _extract_excel_legacy(data, document_id=document_id, file_name=file_name)
    if lower.endswith(".pptx") or "presentation" in mime or "powerpoint" in mime:
        return _extract_pptx(data, document_id=document_id, file_name=file_name)
    if lower.endswith(".rtf") or mime == "application/rtf":
        return _extract_rtf(data, document_id=document_id, file_name=file_name)
    if lower.endswith(_PLAIN_TEXT_EXTENSIONS) or mime.startswith("text/") or mime in (
        "application/json",
        "application/xml",
    ):
        return _extract_plain_text(data, document_id=document_id, file_name=file_name)
    if mime.startswith("image/") or lower.endswith(
        (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tif", ".tiff")
    ):
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="ocr_required",
            warning="图片需文件内容提取后对比",
        )
    return ParsedDocument(
        document_id=document_id,
        file_name=file_name,
        full_text="",
        parse_quality="unsupported",
        warning=f"暂不支持的文件格式: {file_name}",
    )


def _extract_plain_text(
    data: bytes, *, document_id: uuid.UUID, file_name: str
) -> ParsedDocument:
    text = data.decode("utf-8", errors="replace")
    return _parsed_text_doc(
        document_id=document_id,
        file_name=file_name,
        full_text=text,
        parse_quality="text_layer",
    )


def _extract_html(data: bytes, *, document_id: uuid.UUID, file_name: str) -> ParsedDocument:
    try:
        raw = data.decode("utf-8", errors="replace")
        md = html_to_markdown(raw)
        return _parsed_text_doc(
            document_id=document_id,
            file_name=file_name,
            full_text=md or raw,
            parse_quality="html",
        )
    except Exception as e:
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="failed",
            warning=f"HTML 解析失败: {e}",
        )


def _extract_excel(
    data: bytes, *, document_id: uuid.UUID, file_name: str
) -> ParsedDocument:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="unsupported",
            warning="未安装 openpyxl，无法解析 Excel",
        )
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() if c is not None else "" for c in row]
                if any(cells):
                    rows.append("\t".join(cells))
            if rows:
                parts.append(f"## {sheet.title}\n" + "\n".join(rows))
        wb.close()
        full = "\n\n".join(parts).strip()
        return _parsed_text_doc(
            document_id=document_id,
            file_name=file_name,
            full_text=full,
            parse_quality="excel" if full else "failed",
            warning=None if full else "Excel 中未提取到文本",
        )
    except Exception as e:
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="failed",
            warning=f"Excel 解析失败: {e}",
        )


def _extract_excel_legacy(
    data: bytes, *, document_id: uuid.UUID, file_name: str
) -> ParsedDocument:
    try:
        import pandas as pd
    except ImportError:
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="unsupported",
            warning="未安装 pandas，无法解析旧版 .xls",
        )
    try:
        book = pd.read_excel(io.BytesIO(data), sheet_name=None, header=None)
        parts: list[str] = []
        for sheet_name, frame in book.items():
            rows: list[str] = []
            for _, row in frame.iterrows():
                cells = [str(c).strip() if pd.notna(c) else "" for c in row.tolist()]
                if any(cells):
                    rows.append("\t".join(cells))
            if rows:
                parts.append(f"## {sheet_name}\n" + "\n".join(rows))
        full = "\n\n".join(parts).strip()
        return _parsed_text_doc(
            document_id=document_id,
            file_name=file_name,
            full_text=full,
            parse_quality="excel" if full else "failed",
            warning=None if full else "Excel 中未提取到文本",
        )
    except Exception as e:
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="failed",
            warning=f"旧版 Excel 解析失败: {e}",
        )


def _extract_pptx(
    data: bytes, *, document_id: uuid.UUID, file_name: str
) -> ParsedDocument:
    try:
        slides: list[str] = []
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            slide_names = sorted(
                n
                for n in zf.namelist()
                if n.startswith("ppt/slides/slide") and n.endswith(".xml")
            )
            for idx, name in enumerate(slide_names, start=1):
                root = ElementTree.fromstring(zf.read(name))
                texts: list[str] = []
                for node in root.iter():
                    tag = node.tag.rsplit("}", 1)[-1]
                    if tag == "t" and node.text and node.text.strip():
                        texts.append(node.text.strip())
                if texts:
                    slides.append(f"## Slide {idx}\n" + "\n".join(texts))
        full = "\n\n".join(slides).strip()
        return _parsed_text_doc(
            document_id=document_id,
            file_name=file_name,
            full_text=full,
            parse_quality="pptx" if full else "failed",
            warning=None if full else "PPT 中未提取到文本",
        )
    except zipfile.BadZipFile:
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="failed",
            warning="无效的 PPTX 文件",
        )
    except Exception as e:
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="failed",
            warning=f"PPT 解析失败: {e}",
        )


def _extract_rtf(data: bytes, *, document_id: uuid.UUID, file_name: str) -> ParsedDocument:
    try:
        raw = data.decode("utf-8", errors="replace")
        if not raw.lstrip().startswith("{\\rtf"):
            raw = data.decode("latin-1", errors="replace")
        text = re.sub(r"\\[a-z]+\d* ?|\\\{|\\\}|\\'[0-9a-fA-F]{2}", " ", raw)
        text = text.replace("{", " ").replace("}", " ")
        text = unescape(re.sub(r"\s+", " ", text)).strip()
        return _parsed_text_doc(
            document_id=document_id,
            file_name=file_name,
            full_text=text,
            parse_quality="rtf" if text else "failed",
            warning=None if text else "RTF 中未提取到文本",
        )
    except Exception as e:
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="failed",
            warning=f"RTF 解析失败: {e}",
        )


def _extract_pdf(data: bytes, *, document_id: uuid.UUID, file_name: str) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="unsupported",
            warning="未安装 pypdf，无法解析 PDF",
        )

    reader = PdfReader(io.BytesIO(data))
    pages: list[dict] = []
    parts: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        parts.append(text)
        pages.append(
            {
                "page": i,
                "text": text,
                "blocks": [{"text": text, "bbox": None}] if text else [],
            }
        )
    full = "\n\n".join(parts).strip()
    quality = "text_layer" if full else "ocr_required"
    warning = None if full else "PDF 无文本层，后续可进行文件内容提取"
    return ParsedDocument(
        document_id=document_id,
        file_name=file_name,
        full_text=full,
        pages=pages,
        parse_quality=quality,
        warning=warning,
    )


def _extract_docx(
    data: bytes, *, document_id: uuid.UUID, file_name: str
) -> ParsedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="unsupported",
            warning="未安装 python-docx，无法解析 Word 文档",
        )
    try:
        doc = DocxDocument(io.BytesIO(data))
        parts: list[str] = []
        for p in doc.paragraphs:
            line = p.text.strip()
            if not line:
                continue
            style = (getattr(p.style, "name", None) or "").lower()
            if style.startswith("heading"):
                level = 1
                for ch in style.replace("heading", "").strip():
                    if ch.isdigit():
                        level = max(1, min(6, int(ch)))
                        break
                parts.append(f"{'#' * level} {line}")
            else:
                parts.append(line)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        full = "\n\n".join(parts).strip()
        paragraphs = [p for p in parts if p]
        pages = [
            {
                "page": 1,
                "text": full,
                "blocks": [{"text": t, "bbox": None} for t in paragraphs],
            }
        ]
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text=full,
            pages=pages,
            parse_quality="docx",
        )
    except Exception as e:
        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            full_text="",
            parse_quality="failed",
            warning=f"Word 解析失败: {e}",
        )


def split_paragraphs(text: str) -> list[str]:
    chunks = re.split(r"\n\s*\n+", text)
    return [c.strip() for c in chunks if c.strip()]


def _parse_field_query(query: str) -> tuple[list[str], list[tuple[str, str]]]:
    """解析自然语言 + 字段匹配，如「违约金」或「条款:违约」。"""
    fields: list[tuple[str, str]] = []
    terms: list[str] = []
    for part in re.split(r"\s+", query.strip()):
        if not part:
            continue
        if ":" in part or "：" in part:
            sep = ":" if ":" in part else "："
            k, v = part.split(sep, 1)
            k, v = k.strip(), v.strip()
            if k and v:
                fields.append((k.lower(), v.lower()))
                continue
        if len(part) >= 2:
            terms.append(part.lower())
    if not terms and not fields:
        terms = [query.strip().lower()]
    return terms, fields


def local_search(
    parsed_docs: list[ParsedDocument],
    query: str,
    *,
    limit: int = 20,
    field_match: bool = True,
) -> list[dict]:
    """关键词 + 可选字段匹配检索（KnowFlow 不可用时的回退）。"""
    q = query.strip()
    if not q:
        return []
    terms, fields = _parse_field_query(q) if field_match else ([q.lower()], [])
    if not terms and not fields:
        terms = [q.lower()]
    hits: list[dict] = []
    for doc in parsed_docs:
        for page in doc.pages or [{"page": 1, "text": doc.full_text}]:
            page_no = page.get("page", 1)
            for para in split_paragraphs(page.get("text") or ""):
                lower = para.lower()
                score = sum(lower.count(t) for t in terms)
                for fk, fv in fields:
                    if fk in lower and fv in lower:
                        score += 5
                if score <= 0:
                    continue
                hits.append(
                    {
                        "document_id": str(doc.document_id),
                        "snippet": para[:500],
                        "score": float(score),
                        "anchor_json": {"page": page_no, "bbox": None, "kind": "text"},
                        "source": "local",
                    }
                )
        if not doc.pages and doc.full_text:
            for para in split_paragraphs(doc.full_text):
                lower = para.lower()
                score = sum(lower.count(t) for t in terms)
                for fk, fv in fields:
                    if fk in lower and fv in lower:
                        score += 5
                if score <= 0:
                    continue
                hits.append(
                    {
                        "document_id": str(doc.document_id),
                        "snippet": para[:500],
                        "score": float(score),
                        "anchor_json": {"page": 1, "bbox": None, "kind": "text"},
                        "source": "local",
                    }
                )
    hits.sort(key=lambda x: x["score"], reverse=True)
    return hits[:limit]
