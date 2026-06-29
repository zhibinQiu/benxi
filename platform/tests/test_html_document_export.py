"""HTML 导出为 KnowFlow 可索引格式。"""

from unittest.mock import patch

from app.integrations.html_document_export import (
    MIN_ARTICLE_PLAIN_CHARS,
    build_substantive_article_markdown,
    html_body_to_indexable_bytes,
    is_thin_article_content,
    normalize_file_for_knowflow_upload,
    plain_text_char_count,
    resolve_article_html_body,
)


def test_build_substantive_includes_summary_and_body():
    md = build_substantive_article_markdown(
        "测试标题",
        "<p>第一段正文内容。" + "续写。" * 40 + "</p>",
        summary="这是摘要说明",
        link="https://example.com/a",
        source_label="示例站",
    )
    assert "# 测试标题" in md
    assert "摘要" in md
    assert "第一段正文" in md
    assert "## 正文" in md
    assert "https://example.com/a" in md
    assert "示例站" in md
    assert plain_text_char_count(md) >= MIN_ARTICLE_PLAIN_CHARS


def test_resolve_refetches_when_thin():
    thin = "<p>短</p>"
    rich = "<p>" + "长文。" * 80 + "</p>"
    with patch(
        "app.integrations.html_document_export.refetch_article_html",
        return_value=(rich, "新摘要"),
    ):
        html, summary = resolve_article_html_body(thin, summary="旧", link="https://x.com/1")
    assert "长文" in html
    assert summary == "新摘要"


def test_html_body_to_indexable_bytes_produces_markdown():
    body = "<p>正文 " + "内容。" * 60 + "</p>"
    name, content, mime = html_body_to_indexable_bytes(
        "测试标题",
        body,
        summary="摘要",
        link="https://example.com/a",
        allow_refetch=False,
    )
    assert name.endswith(".md")
    assert "text/markdown" in mime
    text = content.decode("utf-8")
    assert not is_thin_article_content(text)
    assert "# 测试标题" in text


def test_normalize_converts_html_upload_to_pdf():
    html = (
        b"<!DOCTYPE html><html><body><p>"
        + ("知识库同步需要有足够长度的中文正文。" * 12).encode()
        + b"</p></body></html>"
    )
    name, content, mime = normalize_file_for_knowflow_upload(
        "article.html", html, "text/html; charset=utf-8", title="文章"
    )
    assert name.endswith(".pdf")
    assert mime == "application/pdf"
    assert content.startswith(b"%PDF")


def test_normalize_converts_thin_markdown_to_pdf_with_link():
    thin_md = b"# t\n\n[link](https://example.com/x)"
    rich = "<p>" + "网页正文。" * 80 + "</p>"
    with patch(
        "app.integrations.html_document_export.resolve_article_html_body",
        return_value=(rich, "sum"),
    ):
        name, content, mime = normalize_file_for_knowflow_upload(
            "t.md",
            thin_md,
            "text/markdown",
            title="文章",
            description="来源：web\n链接：https://example.com/x",
        )
    assert name.endswith(".pdf")
    assert mime == "application/pdf"
    assert content.startswith(b"%PDF")
    assert len(content) > 500


def test_normalize_keeps_pdf_unchanged():
    raw = b"%PDF-1.4"
    name, content, mime = normalize_file_for_knowflow_upload(
        "doc.pdf", raw, "application/pdf"
    )
    assert name == "doc.pdf"
    assert content == raw
    assert mime == "application/pdf"


def test_normalize_converts_plain_text_to_pdf():
    text = ("员工体检通知正文。" * 30).encode("utf-8")
    name, content, mime = normalize_file_for_knowflow_upload(
        "notice.txt", text, "text/plain", title="体检通知"
    )
    assert name.endswith(".pdf")
    assert mime == "application/pdf"
    assert content.startswith(b"%PDF")
    assert len(content) > 500


def test_normalize_keeps_docx_original():
    name, content, mime = normalize_file_for_knowflow_upload(
        "report.docx",
        b"PK fake",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        title="报告",
    )
    assert name == "report.docx"
    assert mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert content == b"PK fake"


def test_normalize_keeps_pptx_original():
    name, content, mime = normalize_file_for_knowflow_upload(
        "slides.pptx",
        b"PK fake pptx",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        title="演示",
    )
    assert name == "slides.pptx"
    assert content == b"PK fake pptx"


def test_normalize_keeps_csv_original():
    csv_bytes = b"col1,col2\na,b\n"
    name, content, mime = normalize_file_for_knowflow_upload(
        "data.csv",
        csv_bytes,
        "text/csv",
        title="数据",
    )
    assert name == "data.csv"
    assert content == csv_bytes
    assert mime == "text/csv"


def test_normalize_keeps_xlsx_original():
    name, content, mime = normalize_file_for_knowflow_upload(
        "sheet.xlsx",
        b"PK fake xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        title="表格",
    )
    assert name == "sheet.xlsx"
    assert content == b"PK fake xlsx"


def test_convert_docx_to_pdf_for_citation_real():
    import io

    from docx import Document

    from app.integrations.html_document_export import convert_file_bytes_to_pdf_for_citation

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("建设生态强企服务碳达峰行动工作计划。" * 5)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "目标"
    table.cell(1, 0).text = "碳排放"
    table.cell(1, 1).text = "下降 10%"
    doc.save(buf)
    result = convert_file_bytes_to_pdf_for_citation(
        "计划.docx",
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        title="工作计划",
    )
    assert result is not None
    name, pdf_bytes, mime = result
    assert name.endswith(".pdf")
    assert mime == "application/pdf"
    assert pdf_bytes.startswith(b"%PDF")
    assert len(pdf_bytes) > 200


def test_knowflow_copy_lacks_page_snapshots_for_block_pdf():
    from app.integrations.html_document_export import knowflow_copy_lacks_page_snapshots

    assert knowflow_copy_lacks_page_snapshots(
        "report.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "report.pdf",
    )
    assert not knowflow_copy_lacks_page_snapshots(
        "report.pdf",
        "application/pdf",
        "report.pdf",
    )
